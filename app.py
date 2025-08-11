import streamlit as st
from docx import Document
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
import json
import re
import os

CHROMA_DIR = "vector_store"
embedding_fn = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
vectorstore = Chroma(persist_directory=CHROMA_DIR, embedding_function=embedding_fn)

REQUIRED_DOCS_INCORP = [
    "Articles of Association",
    "Memorandum of Association",
    "Board Resolution",
    "Shareholder Resolution",
    "Register of Members and Directors"
]

def detect_document_type(text):
    text_lower = text.lower()
    if "articles of association" in text_lower:
        return "Articles of Association"
    elif "memorandum of association" in text_lower:
        return "Memorandum of Association"
    elif "board resolution" in text_lower:
        return "Board Resolution"
    elif "shareholder resolution" in text_lower:
        return "Shareholder Resolution"
    elif "register of members" in text_lower or "register of directors" in text_lower:
        return "Register of Members and Directors"
    return "Unknown"

def detect_red_flags(text):
    """Simple keyword-based red flag detection."""
    flags = []
    if re.search(r"UAE Federal Courts", text, re.IGNORECASE):
        flags.append({
            "issue": "Incorrect jurisdiction reference",
            "suggestion": "Replace with 'ADGM Courts' as per ADGM Companies Regulations."
        })
    if not re.search(r"Signed by|Signature", text, re.IGNORECASE):
        flags.append({
            "issue": "Missing signatory section",
            "suggestion": "Include a signatory section with name, designation, and date."
        })
    if re.search(r"\bmay\b", text, re.IGNORECASE):
        flags.append({
            "issue": "Ambiguous language ('may')",
            "suggestion": "Consider using 'shall' for legally binding obligations."
        })
    return flags

def add_comments_to_docx(doc, issues):
    """Insert comments into a DOCX file for each issue."""
    for issue in issues:
        comment_text = f"{issue['issue']} — Suggestion: {issue['suggestion']}"
        # Add as a new paragraph with highlighting (since python-docx doesn't natively support Word comments)
        p = doc.add_paragraph()
        run = p.add_run(comment_text)
        run.font.highlight_color = 2  # Yellow highlight
    return doc

st.set_page_config(page_title="ADGM Corporate Agent", layout="wide")
st.title("🏛️ ADGM Corporate Agent — Enhanced")
st.write("Upload your legal documents for compliance review against ADGM rules with inline comments.")

uploaded_files = st.file_uploader("Upload .docx files", type=["docx"], accept_multiple_files=True)

if uploaded_files:
    detected_docs = []
    results_summary = []
    reviewed_files = []

    for file in uploaded_files:
        doc = Document(file)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

        doc_type = detect_document_type(text)
        detected_docs.append(doc_type)

        # RAG search for related rules
        query = f"Check ADGM compliance for: {text[:1000]}"
        similar_docs = vectorstore.similarity_search(query, k=3)

        rag_issues = [{
            "source_url": match.metadata.get("source_url"),
            "category": match.metadata.get("category"),
            "document_type": match.metadata.get("document_type"),
            "reference_excerpt": match.page_content[:300] + "..."
        } for match in similar_docs]

        # Red flag detection
        red_flags = detect_red_flags(text)

        # Combine issues for commenting
        all_issues = red_flags.copy()

        results_summary.append({
            "uploaded_file": file.name,
            "detected_document_type": doc_type,
            "rag_findings": rag_issues,
            "red_flags": red_flags
        })

        # Create reviewed doc with inline comments
        reviewed_doc = add_comments_to_docx(doc, red_flags)
        reviewed_filename = f"reviewed_{file.name}"
        reviewed_path = os.path.join("reviewed_docs", reviewed_filename)
        os.makedirs("reviewed_docs", exist_ok=True)
        reviewed_doc.save(reviewed_path)
        reviewed_files.append(reviewed_path)

    # Checklist Verification
    uploaded_set = set(d for d in detected_docs if d != "Unknown")
    missing_docs = [doc for doc in REQUIRED_DOCS_INCORP if doc not in uploaded_set]

    st.subheader("📋 Checklist Verification")
    st.write(
        f"It appears that you're trying to incorporate a company in ADGM. "
        f"Based on our reference list, you have uploaded {len(uploaded_set)} out of {len(REQUIRED_DOCS_INCORP)} required documents."
    )
    if missing_docs:
        st.error(f"Missing document(s): {', '.join(missing_docs)}")
    else:
        st.success("All required documents are present.")

    # JSON Summary
    st.subheader("📄 Review Summary (JSON)")
    st.json({
        "process": "Company Incorporation",
        "documents_uploaded": len(uploaded_set),
        "required_documents": len(REQUIRED_DOCS_INCORP),
        "missing_documents": missing_docs,
        "detailed_results": results_summary
    })

    # Download reviewed DOCX files
    for reviewed_path in reviewed_files:
        with open(reviewed_path, "rb") as f:
            st.download_button(
                label=f"Download Reviewed: {os.path.basename(reviewed_path)}",
                data=f,
                file_name=os.path.basename(reviewed_path),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # Download JSON
    json_bytes = json.dumps({
        "process": "Company Incorporation",
        "documents_uploaded": len(uploaded_set),
        "required_documents": len(REQUIRED_DOCS_INCORP),
        "missing_documents": missing_docs,
        "detailed_results": results_summary
    }, indent=2).encode("utf-8")

    st.download_button(
        label="Download JSON Summary",
        data=json_bytes,
        file_name="adgm_review_summary.json",
        mime="application/json"
    )
