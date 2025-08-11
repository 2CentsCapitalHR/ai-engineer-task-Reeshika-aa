import os
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import mimetypes
import pdfplumber
from docx import Document as DocxReader
import json
from groq import Groq
from dotenv import load_dotenv

# ====== LOAD ENV & GROQ SETUP ======
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise ValueError("GROQ_API_KEY not set in environment or .env file")

client = Groq(api_key=GROQ_API_KEY)

# ====== CONFIG ======
DOCX_SOURCE_FILE = r"C:\\Users\\every\\Desktop\\2cents\\Data Sources.docx"
RAW_DIR = "data/raw"
TEXT_DIR = "data/text"
META_FILE = "data/metadata.json"  # Stores category/doc_type mapping
REVIEW_DIR = "data/reviewed"

os.makedirs(RAW_DIR, exist_ok=True)
os.makedirs(TEXT_DIR, exist_ok=True)
os.makedirs(REVIEW_DIR, exist_ok=True)

# ====== FUNCTION: Groq compliance review with chunking ======
def detect_red_flags_with_groq(text, max_chars=5000):
    """Send document text to Groq for ADGM compliance review in chunks if too long."""
    issues_all = []
    
    # Split into chunks
    chunks = [text[i:i+max_chars] for i in range(0, len(text), max_chars)]
    
    for idx, chunk in enumerate(chunks, start=1):
        prompt = f"""
        You are an ADGM corporate compliance expert.
        Review the text for:
        - Wrong jurisdiction references
        - Missing signature blocks
        - Ambiguous or non-binding clauses
        - Non-compliance with ADGM company setup requirements
        Provide JSON array:
        [
            {{
                "issue": "...",
                "suggestion": "...",
                "reference": "ADGM Companies Regulations / Rule Name"
            }}
        ]
        Text (part {idx} of {len(chunks)}):
        {chunk}
        """

        resp = client.chat.completions.create(
            model="llama3-8b-8192",  # working model
            messages=[
                {"role": "system", "content": "You are an ADGM compliance assistant."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=800
        )

        try:
            part_issues = json.loads(resp.choices[0].message.content)
            if isinstance(part_issues, list):
                issues_all.extend(part_issues)
        except Exception:
            issues_all.append({
                "issue": "Model output not JSON",
                "suggestion": resp.choices[0].message.content
            })

    return issues_all

# ====== 1. Read the table from Data Sources.docx ======
doc = Document(DOCX_SOURCE_FILE)
entries = []

for table in doc.tables:
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) >= 3 and "http" in cells[2]:
            entries.append({
                "category": cells[0],
                "document_type": cells[1],
                "url": cells[2]
            })

print(f"[INFO] Found {len(entries)} entries with URLs")

# ====== 2. Helper functions ======
def download_file(url, save_dir):
    """Download file from URL and return local path + mime type."""
    try:
        r = requests.get(url, stream=True)
        r.raise_for_status()
        content_type = r.headers.get("Content-Type", "")
        ext = mimetypes.guess_extension(content_type.split(";")[0]) or url.split("?")[0].split(".")[-1]
        if not ext.startswith("."):
            ext = "." + ext
        filename = url.split("/")[-1].split("?")[0]
        if not filename.endswith(ext):
            filename += ext
        filepath = os.path.join(save_dir, filename)
        with open(filepath, "wb") as f:
            f.write(r.content)
        return filepath, content_type
    except Exception as e:
        print(f"[ERROR] Failed to download {url}: {e}")
        return None, None

def extract_text_pdf(filepath):
    text = ""
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_text_docx(filepath):
    doc = DocxReader(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)

def extract_text_html(url):
    r = requests.get(url)
    soup = BeautifulSoup(r.text, "html.parser")
    for script in soup(["script", "style"]):
        script.decompose()
    return " ".join(soup.stripped_strings)

def add_comments_to_docx(original_path, issues, output_path):
    """Create a reviewed DOCX with highlights and comments."""
    doc = DocxReader(original_path)
    for para in doc.paragraphs:
        for issue in issues:
            if issue["issue"] and issue["issue"].lower() in para.text.lower():
                # Highlight paragraph
                for run in para.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                # Add comment
                para.add_comment(
                    f"Issue: {issue['issue']}\nSuggestion: {issue['suggestion']}\nReference: {issue['reference']}"
                )
    doc.save(output_path)

# ====== 3. Download, extract, review, save ======
metadata_records = []

for entry in entries:
    category = entry["category"]
    doc_type = entry["document_type"]
    url = entry["url"]

    print(f"[INFO] Processing: {doc_type} | {url}")
    file_path, ctype = download_file(url, RAW_DIR)
    if not file_path:
        continue

    # Extract text
    extracted_text = ""
    if "pdf" in ctype or file_path.lower().endswith(".pdf"):
        extracted_text = extract_text_pdf(file_path)
    elif "word" in ctype or file_path.lower().endswith(".docx"):
        extracted_text = extract_text_docx(file_path)
    elif "html" in ctype or file_path.lower().endswith((".html", ".htm")):
        extracted_text = extract_text_html(url)
    else:
        print(f"[WARN] Unknown type for {url} ({ctype}), skipping text extraction.")

    if extracted_text:
        # Save raw text
        safe_name = os.path.splitext(os.path.basename(file_path))[0]
        text_path = os.path.join(TEXT_DIR, f"{safe_name}.txt")
        with open(text_path, "w", encoding="utf-8") as f:
            f.write(extracted_text)

        # ===== Run Groq review =====
        issues = detect_red_flags_with_groq(extracted_text)

        # ===== Save reviewed DOCX =====
        reviewed_path = os.path.join(REVIEW_DIR, f"{safe_name}_reviewed.docx")
        if file_path.lower().endswith(".docx"):
            add_comments_to_docx(file_path, issues, reviewed_path)

        # ===== Save metadata =====
        metadata_records.append({
            "category": category,
            "document_type": doc_type,
            "source_url": url,
            "raw_file": file_path,
            "text_file": text_path,
            "reviewed_docx": reviewed_path if os.path.exists(reviewed_path) else None,
            "issues_found": issues
        })

# ====== 4. Save metadata JSON ======
with open(META_FILE, "w", encoding="utf-8") as f:
    json.dump(metadata_records, f, indent=2)

print(f"[DONE] Processed {len(metadata_records)} documents.")
print(f"[INFO] Metadata saved to {META_FILE}")
